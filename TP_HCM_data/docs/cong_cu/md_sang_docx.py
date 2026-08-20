# -*- coding: utf-8 -*-
"""Chuyển Markdown sang .docx — dùng khi máy không có pandoc.

Chạy:  py -3.11 md_sang_docx.py TECH_DOC.md
       py -3.11 md_sang_docx.py TECH_DOC.md ten_khac.docx

Hỗ trợ: tiêu đề · bảng · danh sách · khối code · trích dẫn · đường kẻ ngang ·
        định dạng inline (**đậm**, *nghiêng*, `code`) · **nhúng ảnh** ·
        **công thức LaTeX** ($...$ và $$...$$) dịch sang Unicode.
Không hỗ trợ: HTML thô.

Nhúng ảnh — hai cú pháp:
    ![chú thích](hinh_anh/TEN.png)      markdown chuẩn
    🖼️ `TEN_HINH.png` · `TEN_KHAC.png`  quy ước của dự án, tự tìm trong hinh_anh/
"""
import re
import sys
import unicodedata
from pathlib import Path

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

import latex_sang_chu

THU_MUC_HINH = "hinh_anh"      # tìm ảnh ở đây, tương đối với file .md
RONG_ANH = Cm(16.5)

# ── kiểu chữ ─────────────────────────────────────────────────────────────
CHU = "Times New Roman"        # toàn tài liệu
CHU_MA = "Consolas"            # RIÊNG khối code: có sơ đồ ASCII, đổi font là vỡ hàng
CO_THAN = Pt(12.5)               # cỡ chữ thân bài
CO_BANG = Pt(10.5)               # bảng: 13pt thì bảng 6 cột vỡ hết, hạ xuống 11
CO_MA = Pt(9)
CO_TIEU_DE = {1: Pt(17), 2: Pt(14.5), 3: Pt(13), 4: Pt(12.5)}

INK = RGBColor(0x22, 0x22, 0x22)
MUT = RGBColor(0x66, 0x66, 0x66)
ACCENT = RGBColor(0x00, 0x72, 0xB2)


def dat_font(doi_tuong, ten=CHU):
    """Đặt font cho run/style — phải ghi cả 4 khe rFonts.

    python-docx chỉ set `w:ascii`. Word và WPS đọc thêm `hAnsi`, `eastAsia`, `cs`;
    thiếu chúng thì phần dấu tiếng Việt bị thay bằng font mặc định và ra lỗi
    "Font Missing".
    """
    font = doi_tuong.font if hasattr(doi_tuong, "font") else doi_tuong
    font.name = ten
    rPr = font.element.get_or_add_rPr()
    rf = rPr.find(qn("w:rFonts"))
    if rf is None:
        rf = OxmlElement("w:rFonts")
        rPr.append(rf)
    for khe in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(khe), ten)


# ── mục lục bấm được: bookmark ở tiêu đề + liên kết nội bộ ở mục lục ────
_SO_BOOKMARK = [0]


def slug(s):
    """Tên bookmark hợp lệ của Word: bỏ dấu, chỉ chữ/số/gạch dưới, ≤40 ký tự."""
    s = unicodedata.normalize("NFD", s)
    s = "".join(c for c in s if unicodedata.category(c) != "Mn").replace("đ", "d").replace("Đ", "D")
    s = re.sub(r"[^0-9A-Za-z]+", "_", s).strip("_")
    return ("m_" + s)[:40] or "m_x"


def dat_bookmark(para, ten):
    """Cắm bookmark bao quanh paragraph để liên kết nội bộ nhảy tới được."""
    _SO_BOOKMARK[0] += 1
    bid = str(_SO_BOOKMARK[0])
    b1 = OxmlElement("w:bookmarkStart"); b1.set(qn("w:id"), bid); b1.set(qn("w:name"), ten)
    b2 = OxmlElement("w:bookmarkEnd");   b2.set(qn("w:id"), bid)
    para._p.insert(0, b1)
    para._p.append(b2)


def lien_ket_trong(para, chu, neo, dam=False):
    """Chèn liên kết nội bộ (nhảy tới bookmark `neo`) — bấm được trong Word."""
    h = OxmlElement("w:hyperlink"); h.set(qn("w:anchor"), neo)
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rf = OxmlElement("w:rFonts")
    for khe in ("w:ascii", "w:hAnsi", "w:eastAsia", "w:cs"):
        rf.set(qn(khe), CHU)
    rPr.append(rf)
    if dam:
        rPr.append(OxmlElement("w:b"))
    col = OxmlElement("w:color"); col.set(qn("w:val"), "0072B2"); rPr.append(col)
    u = OxmlElement("w:u"); u.set(qn("w:val"), "none"); rPr.append(u)
    r.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = chu
    r.append(t)
    h.append(r)
    para._p.append(h)


# ── công thức ────────────────────────────────────────────────────────────
def viet_cong_thuc(para, latex, co=None, dam=False):
    """Ghi công thức LaTeX vào paragraph — chỉ số dưới/trên là run thật."""
    for chu, kieu in latex_sang_chu.dich(latex):
        r = para.add_run(chu)
        r.italic = True
        r.bold = True if dam else None
        dat_font(r)
        if co:
            r.font.size = co
        if kieu == "sub":
            r.font.subscript = True
        elif kieu == "sup":
            r.font.superscript = True
    return para


# ── inline: **đậm** · *nghiêng* · `code` · $công thức$ ───────────────────
TOKEN = re.compile(
    r"((?<!!)\[[^\]]+\]\(#[^)]+\)"          # [chữ](#neo) — liên kết nội bộ, đặt trước
    r"|\$\$.+?\$\$|\$[^$\n]+\$|\*\*.+?\*\*"
    r"|(?<!\*)\*(?!\*).+?(?<!\*)\*(?!\*)|`[^`]+`)")
NEO = re.compile(r"^\[([^\]]+)\]\(#([^)]+)\)$")


def dat_dam_nghieng(run, dam, nghieng):
    """Chỉ ép đậm/nghiêng khi cần; None để run kế thừa style.

    Ép r.bold = False sẽ ĐÈ lên style Heading (vốn bold=True) và làm mọi đầu mục
    mất in đậm — run luôn thắng style trong OOXML.
    """
    run.bold = True if dam else None
    run.italic = True if nghieng else None


def viet_inline(para, text, dam=False, nghieng=False):
    """Ghi text vào paragraph, tách các đoạn định dạng inline.

    Đệ quy vào bên trong **đậm** và *nghiêng*: nếu không, một công thức nằm lồng
    trong cụm in đậm sẽ bị nuốt nguyên và lòi ra `$t - \\Delta$` như văn bản thô.
    """
    for phan in TOKEN.split(text):
        if not phan:
            continue
        m_neo = NEO.match(phan)
        if m_neo:
            lien_ket_trong(para, m_neo.group(1), m_neo.group(2), dam=dam)
        elif phan.startswith("$$") and phan.endswith("$$") and len(phan) > 4:
            viet_cong_thuc(para, phan[2:-2], dam=dam)
        elif phan.startswith("$") and phan.endswith("$") and len(phan) > 2:
            viet_cong_thuc(para, phan[1:-1], dam=dam)
        elif phan.startswith("**") and phan.endswith("**"):
            viet_inline(para, phan[2:-2], True, nghieng)
        elif phan.startswith("`") and phan.endswith("`"):
            r = para.add_run(phan[1:-1])
            dat_dam_nghieng(r, dam, nghieng)
            dat_font(r)
            r.font.color.rgb = ACCENT
        elif phan.startswith("*") and phan.endswith("*"):
            viet_inline(para, phan[1:-1], dam, True)
        else:
            r = para.add_run(phan)
            dat_dam_nghieng(r, dam, nghieng)
            dat_font(r)


def vien_bang(table):
    """Kẻ viền mảnh cho bảng — python-docx không có API sẵn."""
    tbl = table._tbl
    pr = tbl.tblPr
    borders = OxmlElement("w:tblBorders")
    for canh in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{canh}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), "BFBFBF")
        borders.append(e)
    pr.append(borders)


def tach_hang(dong):
    """'| a | b |' -> ['a', 'b']"""
    return [o.strip() for o in dong.strip().strip("|").split("|")]


# ── ảnh ──────────────────────────────────────────────────────────────────
MD_ANH  = re.compile(r"!\[([^\]]*)\]\(([^)]+)\)")
TEN_ANH = re.compile(r"`([^`]+\.(?:png|jpg|jpeg))`", re.I)
# quy uoc cua du an (bao cao tuan 2, 3): {{IMG:ten.png|chu thich}}
IMG_DA  = re.compile(r"\{\{IMG:([^|}]+)(?:\|([^}]*))?\}\}")


def tim_anh(ten, goc: Path):
    """Tim file anh: duong dan truc tiep, roi hinh_anh/ canh file md, roi len 2 cap.

    Tim len cap tren de file md nam trong thu muc con (vd docs/bao_cao_tuan/) van
    dung chung kho anh o docs/hinh_anh/.
    """
    for thu_muc in (goc, goc.parent, goc.parent.parent):
        for ung_vien in (thu_muc / ten, thu_muc / THU_MUC_HINH / ten,
                         thu_muc / THU_MUC_HINH / f"{ten}.png"):
            if ung_vien.exists():
                return ung_vien
    return None


def chen_anh(doc, duong_dan: Path, chu_thich=""):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.add_run().add_picture(str(duong_dan), width=RONG_ANH)
    if chu_thich:
        c = doc.add_paragraph()
        c.alignment = WD_ALIGN_PARAGRAPH.CENTER
        c.paragraph_format.space_after = Pt(12)
        viet_inline(c, chu_thich)
        for r in c.runs:
            r.font.size = Pt(9)
            if not r.bold:
                r.italic = True
            r.font.color.rgb = MUT


def xu_ly_anh(doc, dong, goc: Path):
    """Chen anh neu dong co cu phap anh. Tra ve so anh da chen."""
    n = 0
    for ten, chu_thich in IMG_DA.findall(dong):          # {{IMG:ten.png|chú thích}}
        f = tim_anh(ten.strip(), goc)
        if f:
            chen_anh(doc, f, (chu_thich or "").strip()); n += 1
    for chu_thich, dd in MD_ANH.findall(dong):           # ![chú thích](đường dẫn)
        f = tim_anh(dd, goc)
        if f:
            chen_anh(doc, f, chu_thich); n += 1
    if n == 0 and "🖼" in dong:                           # 🖼️ `ten.png`
        for ten in TEN_ANH.findall(dong):
            f = tim_anh(ten, goc)
            if f:
                chen_anh(doc, f, Path(ten).stem); n += 1
    return n


def la_dong_ngan_cach(dong):
    """Dòng '|---|---:|' của bảng markdown."""
    return bool(re.fullmatch(r"\|[\s:|-]+\|", dong.strip()))


def can_le(spec):
    if spec.endswith(":") and spec.startswith(":"):
        return WD_ALIGN_PARAGRAPH.CENTER
    if spec.endswith(":"):
        return WD_ALIGN_PARAGRAPH.RIGHT
    return WD_ALIGN_PARAGRAPH.LEFT


def chuyen(md_path: Path, docx_path: Path):
    dong = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    st = doc.styles["Normal"]
    dat_font(st)
    st.font.size = CO_THAN
    st.font.color.rgb = INK
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.15

    # Heading 1..4 mac dinh la Calibri Light mau xanh — phai ep tung style,
    # neu khong Word van hien "Calibri (Heading)" nhu anh chup cua WPS.
    for cap, co in CO_TIEU_DE.items():
        try:
            h = doc.styles[f"Heading {cap}"]
        except KeyError:
            continue
        dat_font(h)
        h.font.size = co
        h.font.bold = True
        h.font.color.rgb = ACCENT if cap == 1 else INK
        h.paragraph_format.space_before = Pt(14 if cap <= 2 else 10)
        h.paragraph_format.space_after = Pt(5)
        h.paragraph_format.keep_with_next = True
    for ten in ("List Bullet", "List Number", "Caption"):
        try:
            dat_font(doc.styles[ten])
            doc.styles[ten].font.size = CO_THAN
        except KeyError:
            pass

    for khu in doc.sections:
        khu.top_margin = khu.bottom_margin = Cm(2)
        khu.left_margin = khu.right_margin = Cm(2.2)

    goc = md_path.parent
    i, n_bang, n_code, n_anh, n_ct = 0, 0, 0, 0, 0
    while i < len(dong):
        l = dong[i]

        # ── ảnh ──────────────────────────────────────────────────────────
        if "🖼" in l or MD_ANH.search(l) or IMG_DA.search(l):
            them = xu_ly_anh(doc, l, goc)
            if them:
                n_anh += them
                i += 1
                continue
            # khong tim thay file -> de roi xuong, in ra nhu van ban thuong

        # ── công thức đứng riêng: $$...$$ (co thể trải nhiều dòng) ───────
        if l.strip().startswith("$$"):
            buf = [l.strip()[2:]]
            if not (buf[0].endswith("$$") and buf[0].strip()):
                i += 1
                while i < len(dong) and "$$" not in dong[i]:
                    buf.append(dong[i].strip()); i += 1
                if i < len(dong):
                    buf.append(dong[i].split("$$")[0]); i += 1
            else:
                buf[0] = buf[0][:-2]
                i += 1
            ct = " ".join(x for x in buf if x.strip())
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.paragraph_format.space_before = Pt(8)
            p.paragraph_format.space_after = Pt(8)
            viet_cong_thuc(p, ct, co=Pt(13.5))
            n_ct += 1
            continue

        # ── khối code ────────────────────────────────────────────────────
        if l.startswith("```"):
            i += 1
            buf = []
            while i < len(dong) and not dong[i].startswith("```"):
                buf.append(dong[i]); i += 1
            i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(14)
            p.paragraph_format.space_before = Pt(4)
            r = p.add_run("\n".join(buf))
            dat_font(r, CHU_MA); r.font.size = CO_MA
            r.font.color.rgb = INK
            n_code += 1
            continue

        # ── bảng ─────────────────────────────────────────────────────────
        if l.strip().startswith("|") and i + 1 < len(dong) and la_dong_ngan_cach(dong[i + 1]):
            tieu_de = tach_hang(l)
            canh = [can_le(s) for s in tach_hang(dong[i + 1])]
            i += 2
            than = []
            while i < len(dong) and dong[i].strip().startswith("|"):
                than.append(tach_hang(dong[i])); i += 1

            t = doc.add_table(rows=1, cols=len(tieu_de))
            t.style = "Table Grid"
            t.alignment = WD_TABLE_ALIGNMENT.CENTER
            vien_bang(t)
            for j, o in enumerate(tieu_de):
                cell = t.rows[0].cells[j]
                cell.text = ""
                p = cell.paragraphs[0]
                p.alignment = canh[j] if j < len(canh) else WD_ALIGN_PARAGRAPH.LEFT
                viet_inline(p, o)
                for r in p.runs:
                    r.bold = True
                    r.font.size = CO_BANG
            for hang in than:
                cells = t.add_row().cells
                for j, o in enumerate(hang[:len(tieu_de)]):
                    cells[j].text = ""
                    p = cells[j].paragraphs[0]
                    p.alignment = canh[j] if j < len(canh) else WD_ALIGN_PARAGRAPH.LEFT
                    viet_inline(p, o)
                    for r in p.runs:
                        r.font.size = CO_BANG
            doc.add_paragraph()
            n_bang += 1
            continue

        # ── đường kẻ ngang ───────────────────────────────────────────────
        if l.strip() in ("---", "***", "___"):
            p = doc.add_paragraph()
            pr = p._p.get_or_add_pPr()
            b = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single"); bot.set(qn("w:sz"), "6")
            bot.set(qn("w:color"), "BFBFBF")
            b.append(bot); pr.append(b)
            i += 1
            continue

        # ── tiêu đề ──────────────────────────────────────────────────────
        m = re.match(r"^(#{1,6})\s+(.*)", l)
        if m:
            cap = len(m.group(1))
            p = doc.add_heading("", level=min(cap, 4))
            viet_inline(p, m.group(2))
            dat_bookmark(p, slug(re.sub(r"[*`]", "", m.group(2))))
            for r in p.runs:
                r.font.color.rgb = ACCENT if cap == 1 else INK
                r.font.size = CO_TIEU_DE.get(min(cap, 4))
            i += 1
            continue

        # ── trích dẫn ────────────────────────────────────────────────────
        if l.startswith(">"):
            buf = []
            while i < len(dong) and dong[i].startswith(">"):
                buf.append(dong[i].lstrip(">").strip()); i += 1
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Pt(18)
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            viet_inline(p, " ".join(x for x in buf if x))
            for r in p.runs:
                r.font.color.rgb = MUT
                r.italic = True
            continue

        # ── danh sách ────────────────────────────────────────────────────
        m = re.match(r"^(\s*)[-*+]\s+(.*)", l)
        if m:
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.left_indent = Pt(18 + 14 * (len(m.group(1)) // 2))
            viet_inline(p, m.group(2))
            i += 1
            continue
        m = re.match(r"^(\s*)\d+\.\s+(.*)", l)
        if m:
            p = doc.add_paragraph(style="List Number")
            p.paragraph_format.left_indent = Pt(18 + 14 * (len(m.group(1)) // 2))
            viet_inline(p, m.group(2))
            i += 1
            continue

        # ── đoạn văn (gộp các dòng liền nhau) ────────────────────────────
        if l.strip():
            buf = [l.strip()]
            i += 1
            while (i < len(dong) and dong[i].strip()
                   and not re.match(r"^(#{1,6}\s|[-*+]\s|\d+\.\s|>|```|\|)", dong[i].strip())
                   and dong[i].strip() not in ("---", "***", "___")):
                buf.append(dong[i].strip()); i += 1
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            viet_inline(p, " ".join(buf))
            continue

        i += 1

    doc.save(docx_path)
    return n_bang, n_code, n_anh, n_ct


if __name__ == "__main__":
    if len(sys.argv) < 2:
        sys.exit("dùng: py -3.11 md_sang_docx.py <file.md> [file.docx]")
    src = Path(sys.argv[1]).resolve()
    dst = Path(sys.argv[2]).resolve() if len(sys.argv) > 2 else src.with_suffix(".docx")
    nb, nc, na, nct = chuyen(src, dst)
    print(f"{src.name} -> {dst.name}")
    print(f"  {nb} bảng · {na} hình · {nc} khối code · {nct} công thức · "
          f"{dst.stat().st_size/1024:.0f} KB")
